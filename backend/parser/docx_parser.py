"""DOCX parsing utilities for the ingestion pipeline."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document

from ingestion.models import DocumentContent


def _normalize_metadata(raw_metadata: Any) -> dict[str, str | int | bool]:
    if raw_metadata is None:
        return {}

    normalized: dict[str, str | int | bool] = {}
    document_properties = getattr(raw_metadata, "__dict__", None)
    if isinstance(document_properties, dict):
        for key, value in document_properties.items():
            if value is None:
                continue
            if isinstance(value, (str, int, bool, float)):
                normalized[key] = value
            else:
                normalized[key] = str(value)
    else:
        for key in dir(raw_metadata):
            if key.startswith("_"):
                continue
            try:
                value = getattr(raw_metadata, key)
            except Exception:
                continue
            if value is None or callable(value):
                continue
            if isinstance(value, (str, int, bool, float)):
                normalized[key] = value
            else:
                normalized[key] = str(value)
    return normalized


def _paragraph_text(paragraph) -> str:
    return paragraph.text.strip()


def _table_text(table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def parse_docx(*, docx_bytes: bytes, filename: str) -> DocumentContent:
    """Parse DOCX bytes into a structured DocumentContent object.

    DOCX files do not expose fixed physical pages. For non-paginated document
    formats like DOCX, the returned `pages` tuple represents structural blocks
    such as headings, paragraphs, and tables. In that case, `page_count` is the
    number of structural blocks preserved, not a physical page count.
    """
    if not docx_bytes:
        raise ValueError("docx_bytes must not be empty")
    if not filename:
        raise ValueError("filename must not be empty")

    try:
        document = Document(BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError("Unable to open DOCX bytes") from exc

    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = _paragraph_text(paragraph)
        if text:
            paragraphs.append(text)

    for table in document.tables:
        table_text = _table_text(table)
        if table_text:
            paragraphs.append(table_text)

    text = "\n\n".join(paragraphs)
    metadata = _normalize_metadata(document.core_properties)

    return DocumentContent(
        filename=filename,
        text=text,
        pages=tuple(paragraphs),
        page_count=len(paragraphs),
        metadata=metadata,
    )
