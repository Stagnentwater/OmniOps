"""Unit tests for the DOCX parser."""

import unittest
from io import BytesIO

from docx import Document

from parser.docx_parser import parse_docx


def make_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("This is the first paragraph.")
    document.add_paragraph("This is the second paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Value A"
    table.cell(1, 1).text = "Value B"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class TestDocxParser(unittest.TestCase):
    def test_parse_docx_extracts_text_and_tables(self):
        docx_bytes = make_docx_bytes()

        result = parse_docx(docx_bytes=docx_bytes, filename="document.docx")

        self.assertEqual(result.filename, "document.docx")
        self.assertIn("Project Overview", result.text)
        self.assertIn("This is the first paragraph.", result.text)
        self.assertIn("Header 1 | Header 2", result.text)
        self.assertEqual(result.page_count, 4)
        self.assertEqual(result.pages[0], "Project Overview")
        self.assertEqual(result.pages[1], "This is the first paragraph.")
        self.assertEqual(result.pages[2], "This is the second paragraph.")
        self.assertEqual(result.pages[3], "Header 1 | Header 2\nValue A | Value B")
        self.assertTrue(result.metadata)

    def test_parse_docx_rejects_empty_bytes(self):
        with self.assertRaises(ValueError):
            parse_docx(docx_bytes=b"", filename="empty.docx")

    def test_parse_docx_rejects_empty_filename(self):
        docx_bytes = make_docx_bytes()

        with self.assertRaises(ValueError):
            parse_docx(docx_bytes=docx_bytes, filename="")


if __name__ == "__main__":
    unittest.main()
