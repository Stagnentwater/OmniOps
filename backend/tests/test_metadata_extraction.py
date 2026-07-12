"""Unit tests for metadata extraction."""

import unittest
from io import BytesIO

import fitz
from docx import Document

from ingestion.metadata import (
    DocumentMetadata,
    extract_metadata,
    _detect_file_type,
    _count_words,
    _count_characters,
    _estimate_reading_time,
    _compute_file_hash,
    _detect_has_tables,
    _detect_has_images,
)
from parser.csv_parser import parse_csv
from parser.docx_parser import parse_docx
from parser.pdf_parser import parse_pdf


def make_pdf_bytes(text: str = "Hello OmniOps") -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    document.set_metadata({"title": "Test PDF", "author": "OmniOps"})
    pdf_bytes = document.write()
    document.close()
    return pdf_bytes


def make_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Test DOCX", level=1)
    document.add_paragraph("This is test content.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_csv_bytes() -> bytes:
    csv_content = """Name,Value
Test Asset,100"""
    return csv_content.encode("utf-8")


class TestMetadataHelpers(unittest.TestCase):
    def test_detect_file_type_pdf(self):
        self.assertEqual(_detect_file_type("document.pdf"), "pdf")

    def test_detect_file_type_docx(self):
        self.assertEqual(_detect_file_type("document.docx"), "docx")

    def test_detect_file_type_csv(self):
        self.assertEqual(_detect_file_type("data.csv"), "csv")

    def test_detect_file_type_unknown(self):
        self.assertEqual(_detect_file_type("document.txt"), "unknown")

    def test_detect_file_type_case_insensitive(self):
        self.assertEqual(_detect_file_type("DOCUMENT.PDF"), "pdf")

    def test_count_words(self):
        self.assertEqual(_count_words("hello world"), 2)
        self.assertEqual(_count_words("one two three four"), 4)
        self.assertEqual(_count_words(""), 0)

    def test_count_characters(self):
        self.assertEqual(_count_characters("hello"), 5)
        self.assertEqual(_count_characters("hello world"), 10)
        self.assertEqual(_count_characters(""), 0)

    def test_estimate_reading_time(self):
        self.assertEqual(_estimate_reading_time(200), 1.0)
        self.assertEqual(_estimate_reading_time(400), 2.0)
        self.assertEqual(_estimate_reading_time(0), 0.0)

    def test_compute_file_hash(self):
        data1 = b"test"
        data2 = b"test"
        self.assertEqual(_compute_file_hash(data1), _compute_file_hash(data2))

    def test_detect_has_tables_csv(self):
        parser_metadata = {"delimiter": ",", "row_count": 5}
        self.assertTrue(_detect_has_tables(parser_metadata))

    def test_detect_has_images_true(self):
        parser_metadata = {"has_images": True}
        self.assertTrue(_detect_has_images(parser_metadata))

    def test_detect_has_images_false(self):
        parser_metadata = {"delimiter": ","}
        self.assertFalse(_detect_has_images(parser_metadata))


class TestMetadataExtraction(unittest.TestCase):
    def test_extract_metadata_from_pdf(self):
        pdf_bytes = make_pdf_bytes("Hello OmniOps test content")
        document_content = parse_pdf(pdf_bytes=pdf_bytes, filename="test.pdf")

        metadata = extract_metadata(
            document_content=document_content,
            file_bytes=pdf_bytes,
            document_id="doc-123",
            filename="test.pdf",
            upload_time="2026-07-12T16:40:00Z",
            storage_uri="s3://bucket/test.pdf",
            ingestion_job_id="job-456",
        )

        self.assertIsInstance(metadata, DocumentMetadata)
        self.assertEqual(metadata.document_id, "doc-123")
        self.assertEqual(metadata.filename, "test.pdf")
        self.assertEqual(metadata.upload_time, "2026-07-12T16:40:00Z")
        self.assertEqual(metadata.storage_uri, "s3://bucket/test.pdf")
        self.assertEqual(metadata.ingestion_job_id, "job-456")
        self.assertEqual(metadata.file_type, "pdf")
        self.assertGreater(metadata.word_count, 0)
        self.assertGreater(metadata.character_count, 0)

    def test_extract_metadata_from_docx(self):
        docx_bytes = make_docx_bytes()
        document_content = parse_docx(docx_bytes=docx_bytes, filename="test.docx")

        metadata = extract_metadata(
            document_content=document_content,
            file_bytes=docx_bytes,
            document_id="doc-789",
            filename="test.docx",
            upload_time="2026-07-12T16:40:00Z",
            storage_uri="s3://bucket/test.docx",
            ingestion_job_id="job-789",
        )

        self.assertEqual(metadata.file_type, "docx")
        self.assertEqual(metadata.document_id, "doc-789")

    def test_extract_metadata_from_csv(self):
        csv_bytes = make_csv_bytes()
        document_content = parse_csv(csv_bytes=csv_bytes, filename="test.csv")

        metadata = extract_metadata(
            document_content=document_content,
            file_bytes=csv_bytes,
            document_id="doc-csv",
            filename="test.csv",
            upload_time="2026-07-12T16:40:00Z",
            storage_uri="s3://bucket/test.csv",
            ingestion_job_id="job-csv",
        )

        self.assertEqual(metadata.file_type, "csv")
        self.assertTrue(metadata.has_tables)

    def test_extract_metadata_file_hash_is_deterministic(self):
        file_bytes = b"test content"

        document_content = parse_pdf(
            pdf_bytes=make_pdf_bytes(), filename="test.pdf"
        )

        metadata1 = extract_metadata(
            document_content=document_content,
            file_bytes=file_bytes,
            document_id="doc-1",
            filename="test.pdf",
            upload_time="2026-07-12T16:40:00Z",
            storage_uri="s3://bucket/test.pdf",
            ingestion_job_id="job-1",
        )

        metadata2 = extract_metadata(
            document_content=document_content,
            file_bytes=file_bytes,
            document_id="doc-2",
            filename="test.pdf",
            upload_time="2026-07-12T16:40:01Z",
            storage_uri="s3://bucket/test2.pdf",
            ingestion_job_id="job-2",
        )

        # File hash should be identical for same file bytes
        self.assertEqual(metadata1.file_hash, metadata2.file_hash)

    def test_extract_metadata_system_fields_passed_through(self):
        pdf_bytes = make_pdf_bytes()
        document_content = parse_pdf(pdf_bytes=pdf_bytes, filename="test.pdf")

        metadata = extract_metadata(
            document_content=document_content,
            file_bytes=pdf_bytes,
            document_id="specific-id-123",
            filename="original-name.pdf",
            upload_time="2026-07-12T12:00:00Z",
            storage_uri="azure://container/path/file",
            ingestion_job_id="job-specific-456",
        )

        # System fields must be passed through unchanged
        self.assertEqual(metadata.document_id, "specific-id-123")
        self.assertEqual(metadata.filename, "original-name.pdf")
        self.assertEqual(metadata.upload_time, "2026-07-12T12:00:00Z")
        self.assertEqual(metadata.storage_uri, "azure://container/path/file")
        self.assertEqual(metadata.ingestion_job_id, "job-specific-456")

    def test_extract_metadata_language_default(self):
        pdf_bytes = make_pdf_bytes()
        document_content = parse_pdf(pdf_bytes=pdf_bytes, filename="test.pdf")

        metadata = extract_metadata(
            document_content=document_content,
            file_bytes=pdf_bytes,
            document_id="doc-1",
            filename="test.pdf",
            upload_time="2026-07-12T16:40:00Z",
            storage_uri="s3://bucket/test.pdf",
            ingestion_job_id="job-1",
        )

        # Language should default to unknown (classification deferred)
        self.assertEqual(metadata.language, "unknown")


if __name__ == "__main__":
    unittest.main()
